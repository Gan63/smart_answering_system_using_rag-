import uuid
import os
from docx import Document
from PIL import Image as PILImage
from io import BytesIO
from database.chroma_client import get_text_collection, get_image_collection
from models.embedding_model import embed_text, embed_image

def extract_text_and_images_from_docx(docx_path):
    text = ""
    image_paths = []
    
    # Create directory for extracted images
    image_output_dir = "data/extracted_images"
    os.makedirs(image_output_dir, exist_ok=True)
    
    doc = Document(docx_path)
    
    # Extract text from paragraphs and tables
    for para in doc.paragraphs:
        text += para.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text += para.text + "\n"
    
    # Extract images (inline shapes)
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref.lower():
            image_data = rel.target_part.blob
            image_ext = os.path.splitext(rel.target_ref)[1] or ".png"
            image_filename = f"docx_image_{uuid.uuid4().hex[:8]}{image_ext}"
            image_path = os.path.join(image_output_dir, image_filename)
            
            try:
                pil_image = PILImage.open(BytesIO(image_data))
                pil_image.save(image_path)
                image_paths.append(image_path)
            except Exception as e:
                print(f"Failed to save image {image_filename}: {e}")
    
    return text.strip(), image_paths

def chunk_text(text, chunk_size=300):
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size):
        chunk = " ".join(words[i:i + chunk_size])
        chunks.append(chunk)
    return chunks

def ingest_docx(docx_path, session_id: str):
    source_name = os.path.basename(docx_path)
    print(f"Ingesting DOCX: {docx_path}")
    text, image_paths = extract_text_and_images_from_docx(docx_path)
    
    if not text and not image_paths:
        print("No content found in DOCX.")
        return
    
    # Ingest text chunks
    if text:
        chunks = chunk_text(text)
        text_collection = get_text_collection()
        for chunk in chunks:
            try:
                embedding = embed_text(chunk)
            except Exception as e:
                print(f"Text embedding failed for DOCX chunk: {str(e)[:100]}")
                continue
            try:
                text_collection.add(
                    ids=[str(uuid.uuid4())],
                    embeddings=[embedding],
                    documents=[chunk],
                    metadatas=[{"source": source_name, "session_id": session_id}]
                )
            except Exception as e:
                print(f"Text add to Chroma failed: {str(e)[:100]}")
        print(f"Stored {len(chunks)} text chunks in ChromaDB")
    
    # Ingest images
    if image_paths:
        source_name = os.path.basename(docx_path)
        image_collection = get_image_collection()
        for image_path in image_paths:
            try:
                embedding = embed_image(image_path)
            except Exception as e:
                print(f"Image embedding failed for {image_path}: {str(e)[:100]}")
                continue
            try:
                image_collection.add(
                    ids=[str(uuid.uuid4())],
                    embeddings=[embedding],
                    documents=[image_path.replace("\\", "/")],
                    metadatas=[{"source": source_name, "session_id": session_id}]
                )
            except Exception as e:
                print(f"Image add to Chroma failed: {str(e)[:100]}")
        print(f"Stored {len(image_paths)} images in ChromaDB")
    
    print("[*] Stats updated from Chroma DB")
    
    from session_store import session_store
    session_store.update_stats(session_id)
    print("[+] Updated session stats from Chroma")
    print("DOCX ingestion complete!")

